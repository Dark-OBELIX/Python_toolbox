import os
from docx import Document
from docx.shared import Pt

def copier_paragraphe(destination, source_paragraphe, garder_mise_en_forme):
    nouveau_paragraphe = destination.add_paragraph()

    for run in source_paragraphe.runs:
        nouveau_run = nouveau_paragraphe.add_run(run.text)
        if garder_mise_en_forme:
            nouveau_run.font.name = run.font.name
            nouveau_run.font.size = run.font.size
            # Ajoutez ici d'autres propriétés de mise en forme selon vos besoins

def fusionner_fichiers_docx(dossier_entree, fichier_sortie, garder_mise_en_forme, inclure_titres):
    # Créer un nouveau document pour la sauvegarde
    sauvegarde = Document()

    # Parcourir tous les fichiers dans le dossier d'entrée
    for idx, fichier in enumerate(os.listdir(dossier_entree)):
        if fichier.endswith(".docx"):
            chemin_fichier = os.path.join(dossier_entree, fichier)

            # Ouvrir chaque fichier DOCX et copier son contenu dans le fichier de sauvegarde
            doc_source = Document(chemin_fichier)

            # Inclure le titre du document original
            if inclure_titres:
                titre_original = os.path.splitext(fichier)[0]  # Utiliser le nom du fichier sans extension comme titre
                sauvegarde.add_heading(titre_original, level=1)

            # Copier le contenu du document
            for paragraphe in doc_source.paragraphs:
                copier_paragraphe(sauvegarde, paragraphe, garder_mise_en_forme)

            # Ajouter une page vierge sauf pour le dernier document
            if idx < len(os.listdir(dossier_entree)) - 1:
                sauvegarde.add_page_break()

    # Enregistrer le document de sauvegarde
    sauvegarde.save(fichier_sortie)
    print(f"La fusion des fichiers DOCX est terminée. La sauvegarde a été créée dans {fichier_sortie}")

if __name__ == "__main__":
    # Demander à l'utilisateur de sélectionner un dossier d'entrée
    dossier_entree = input("Veuillez entrer le chemin du dossier contenant les fichiers DOCX : ")

    # Vérifier si le dossier existe
    if not os.path.isdir(dossier_entree):
        print("Le chemin spécifié n'est pas un dossier valide.")
    else:
        # Demander à l'utilisateur le nom du fichier de sauvegarde
        fichier_sortie = input("Veuillez entrer le nom du fichier de sauvegarde (ex: sauvegarde.docx) : ")

        # Demander à l'utilisateur s'il veut garder la mise en forme
        garder_mise_en_forme = input("Voulez-vous conserver la mise en forme d'origine ? (Oui/Non) : ").lower() == "oui"

        # Demander à l'utilisateur s'il veut inclure les titres des documents sources
        inclure_titres = input("Voulez-vous inclure les titres des documents sources ? (Oui/Non) : ").lower() == "oui"

        # Fusionner les fichiers DOCX et créer le fichier de sauvegarde
        fusionner_fichiers_docx(dossier_entree, fichier_sortie, garder_mise_en_forme, inclure_titres)
